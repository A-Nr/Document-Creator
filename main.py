import os
import sys
from pathlib import Path

from PyQt5 import QtWidgets as qtw
from PyQt5 import QtGui as qtg
from PyQt5 import QtCore as qtc
from win32com import client

from MainWindow import Ui_MainWindow
from section import Section
from docx import Document
from doc import sections


class Engine(qtc.QObject):
    generation_done = qtc.pyqtSignal(bool, str)

    def __init__(self):
        super(Engine, self).__init__()
        self._included_sections = set()

    @qtc.pyqtSlot(Section, bool)
    def update_included_sections(self, section, added):
        if added:
            self._included_sections.add(section)
            print('added', section.value)
        else:
            print('removed', section)
            self._included_sections.remove(section)

    def process_section(self, document: Document, section: dict, level: int):
        text = section['text']
        stype = section['type']

        if stype == 'head':
            p = document.add_heading(text, level=level)
        elif stype == 'numbered':
            p = document.add_paragraph(text)
            p.style = 'Numbered Paragraph'
        elif stype == 'paragraph':
            p = document.add_paragraph(text)
            # paragraph_format = paragraph.paragraph_format
            # paragraph_format.space_after = Pt(18)
        elif stype == 'bullet':
            p = document.add_paragraph(text)
            p.style = 'Bullet Black'
        else:
            print(text)
            print('ERROR')
            exit(1)

        paragraph_format = p.paragraph_format
        paragraph_format.keep_together = True

        document.add_paragraph()

        if 'subs' in section:
            for subsection in section['subs']:
                self.process_section(document, subsection, level + 1)

    @qtc.pyqtSlot(str)
    def generate_document(self, path: str):
        path = Path(path).absolute()

        if getattr(sys, 'frozen', False):
            template_path = os.path.join(sys._MEIPASS, 'template.docx')
        else:
            template_path = 'template.docx'

        d = Document(template_path)
        for section in sections:
            if section.get('enum', None) in self._included_sections:
                self.process_section(d, section, 1)

        error = None
        try:
            d.save(path)
        except PermissionError:
            if path.exists():
                error = f'Cannot replace file: {path.name}.\nPlease close the file in Microsoft Word if it\'s opened then try again.'

        if not error:
            try:
                word = client.DispatchEx("Word.Application")
                doc = word.Documents.Open(str(path))
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
            except Exception as e:
                error = str(e)

        if error:
            self.generation_done.emit(False, error)
        else:
            self.generation_done.emit(True, str(path))


class MainWindow(qtw.QMainWindow):
    section_state_changed = qtc.pyqtSignal(Section, bool)
    ready_to_generate = qtc.pyqtSignal(str)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self.engine: Engine = Engine()
        self.engine_thread = qtc.QThread()

        self.engine.moveToThread(self.engine_thread)

        # connect signals from engine
        self.engine.generation_done.connect(self.show_result)

        # connect signals from mainwindow to engine
        self.section_state_changed.connect(self.engine.update_included_sections)
        self.ready_to_generate.connect(self.engine.generate_document)

        self.engine_thread.start()

        # connect buttons to actions
        self.ui.exitButton.clicked.connect(self.close)
        self.ui.generateButton.clicked.connect(self.generate_report)

        # connect checkboxes to signal
        for checkbox in self.ui.groupBox.findChildren(qtw.QCheckBox):
            checkbox.stateChanged.connect(self.section_checbox_clicked)

        # map checboxes to corresponding sections
        self._map_checkboxes_to_sections()

        self.show()

    def _map_checkboxes_to_sections(self):
        self._sections_mapping = {
            self.ui.summaryCheck: Section.Summary,
            self.ui.legalTitleCheck: Section.LegalTitle,
            self.ui.leaseCheck: Section.Lease,
            self.ui.searchResultsCheck: Section.SearchResult,
            self.ui.additionalInformationCheck: Section.AdditionalInformation,
            self.ui.surveyCheck: Section.Survey,
            self.ui.stampLandTaxCheck: Section.StampLandTax,
            self.ui.mortgageCheck: Section.Mortgage,
            self.ui.documentsCheck: Section.Documents,
            self.ui.exchangeCheck: Section.Exchange,
            self.ui.conclusionCheck: Section.Conclusion
        }

    def section_checbox_clicked(self, state):
        checkbox = self.sender()
        section = self._sections_mapping.get(checkbox)
        if section is None:
            return
        if state == qtc.Qt.Checked:
            self.section_state_changed.emit(section, True)
        else:
            self.section_state_changed.emit(section, False)

    def generate_report(self):
        filename, _ = qtw.QFileDialog.getSaveFileName(self, "Save As", "", "Word Document (*.docx)")
        if not filename:
            return

        self.ui.tabs.setEnabled(False)
        self.ready_to_generate.emit(filename)

    def show_result(self, success, text):
        self.ui.tabs.setEnabled(True)
        if success:
            qtw.QMessageBox.information(
                self,
                'Finished Successfully',
                f'The report is saved at:\n{text}'
            )
        else:
            qtw.QMessageBox.critical(
                self,
                'Error',
                f'The following error happend:\n{text}'
            )


if __name__ == '__main__':
    app = qtw.QApplication(sys.argv)
    mw = MainWindow()
    sys.exit(app.exec())
